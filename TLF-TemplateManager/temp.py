import json
from pathlib import Path
from docx import Document

def load_replacements(replacements_file: str) -> dict:
    with open(replacements_file, 'r') as file:
        return json.load(file)

def prompt_for_replacements(replacements: dict) -> dict:
    updated_replacements = {}
    for key in replacements.keys():
        user_input = input(f"Enter value for '{key}': ")
        updated_replacements[key] = user_input
    return updated_replacements        

def replace_text_in_docx(doc: Document, replacements: dict) -> None:
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # Replace text in tables if needed
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

def process_template(template_file: Path, replacements: dict, output_dir: Path) -> None:
    doc = Document(template_file)
    replace_text_in_docx(doc, replacements)
    output_file = output_dir / template_file.name
    doc.save(output_file)

def main(templates_dir: str, replacements_file: str, output_dir: str) -> None:
    templates_path = Path(templates_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    replacements = load_replacements(replacements_file)
    updated_replacements = prompt_for_replacements(replacements)

    for template_file in templates_path.glob('*.docx'):
        process_template(template_file, updated_replacements, output_path)
        print(f"Processed {template_file.name} and saved to {output_path}")

if __name__ == "__main__":
    templates_dir = input("Enter the path to your templates directory: ")
    replacements_file = input("Enter the path to your replacements.json file: ")
    output_dir = input("Enter the path to your output directory: ")

    main(templates_dir, replacements_file, output_dir)
