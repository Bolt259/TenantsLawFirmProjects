import os
import json
import requests
from msal import ConfidentialClientApplication
from docx import Document
from io import BytesIO
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

CLIENT_ID = os.getenv('CLIENT_ID')
CLIENT_SECRET = os.getenv('CLIENT_SECRET')
TENANT_ID = os.getenv('TENANT_ID')
DRIVE_ID = os.getenv('DRIVE_ID')
ITEM_ID = os.getenv('ITEM_ID')

# FUNCTIONS
##################################################################################################################################################################
def authenticate_graph_api():
    """Authenticate with Microsoft Graph API."""
    app = ConfidentialClientApplication(
        CLIENT_ID,
        authority=f"https://login.microsoftonline.com/{TENANT_ID}",
        client_credential=CLIENT_SECRET,
    )
    result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
    if "access_token" in result:
        return result["access_token"]
    else:
        print(result.get("error"))
        print(result.get("error_description"))
        raise Exception("Authentication failed")

def upload_to_onedrive(file_path, item_id, access_token):
    """Upload a file to OneDrive."""
    upload_endpoint = f"https://graph.microsoft.com/v1.0/drives/{DRIVE_ID}/items/{item_id}/content"
    headers = {
        'Authorization': 'Bearer ' + access_token,
        'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }
    with open(file_path, 'rb') as file:
        response = requests.put(upload_endpoint, headers=headers, data=file)
        response.raise_for_status()
        print("File uploaded successfully")

def publish_word_doc(template_path, replacements, output_path, item_id):
    """Replace placeholders in a Word document and upload it to OneDrive."""
    # Authenticate with Microsoft Graph API
    access_token = authenticate_graph_api()

    # Load and modify the document
    document = replace_placeholders(Document(template_path), replacements)

    # Save the modified document
    document.save(output_path)

    # Upload the modified document to OneDrive
    upload_to_onedrive(output_path, item_id, access_token)
    
# Replace placeholders in the document
def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)
    return doc

# Define replacements dynamically
def get_replacements():
    with open('replacements.json', 'r') as file:
        replacements = json.load(file)
    
    return replacements

def get_file_paths(folder_path):
    """
    Get all file paths from a specified folder, including subfolders.

    :param folder_path: Path to the folder
    :return: List of file paths
    """
    file_paths = []
    
    for root, _, files in os.walk(folder_path):
        for file in files:
            file_paths.append(os.path.join(root, file))
    
    return file_paths


##################################################################################################################################################################

# Initialize the MSAL confidential client
app = ConfidentialClientApplication(
    CLIENT_ID,
    authority=f"https://login.microsoftonline.com/{TENANT_ID}",
    client_credential=CLIENT_SECRET,
)

# Acquire an access token for the Graph API
result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])

if "access_token" in result:
    access_token = result["access_token"]
else:
    print(result.get("error"))
    print(result.get("error_description"))
    exit(1)

# Define the endpoint to get the document content
endpoint = f"https://graph.microsoft.com/v1.0/drives/{DRIVE_ID}/items/{ITEM_ID}/content"
headers = {
    'Authorization': 'Bearer ' + access_token
}

# Step 1: Check if the folder exists
folder_name = "%TestFolder"
search_folder_endpoint = f"https://graph.microsoft.com/v1.0/me/drive/root/children?$filter=name eq '{folder_name}'"

response = requests.get(search_folder_endpoint, headers=headers)
response.raise_for_status()

items = response.json().get('value', [])

if items:
    folder_id = items[0]['id']
    print(f"Folder '{folder_name}' already exists with ID: {folder_id}")
else:
    # Step 2: Create the folder if it doesn't exist
    create_folder_endpoint = "https://graph.microsoft.com/v1.0/me/drive/root/children"
    folder_data = {
        "name": folder_name,
        "folder": {},
        "@microsoft.graph.conflictBehavior": "fail"  # This will fail if the folder already exists
    }

    response = requests.post(create_folder_endpoint, headers=headers, json=folder_data)
    response.raise_for_status()

    folder_info = response.json()
    folder_id = folder_info['id']
    print(f"Created folder with ID: {folder_id}")

# Step 3: Upload multiple files to the created or existing directory

# grab filepaths from templates folder
file_paths = get_file_paths(r"C:\Users\sakul\Desktop\VS\Projects\TLF-TemplateManager\templates")

for file_path in file_paths:
    file_name = os.path.basename(file_path)
    upload_endpoint = f"https://graph.microsoft.com/v1.0/me/drive/items/{folder_id}:/{file_name}:/content"

    with open(file_path, 'rb') as file:
        file_content = file.read()

    upload_response = requests.put(upload_endpoint, headers=headers, data=file_content)
    upload_response.raise_for_status()

    print(f"Uploaded file: {file_name}")

# Repeat the above block to upload more files if needed

# Download the document content
response = requests.get(endpoint, headers=headers)
response.raise_for_status()
document = Document(BytesIO(response.content))

# Get dynamic replacements
replacements = get_replacements()

# Modify the document
modified_document = replace_placeholders(document, replacements)

# Save the modified document to a BytesIO object
output_stream = BytesIO()
modified_document.save(output_stream)
output_stream.seek(0)

# Define the endpoint to upload the modified document
upload_endpoint = f"https://graph.microsoft.com/v1.0/drives/{DRIVE_ID}/items/{ITEM_ID}/content"

# Upload the modified document
upload_headers = {
    'Authorization': 'Bearer ' + access_token,
    'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
}
upload_response = requests.put(upload_endpoint, headers=upload_headers, data=output_stream)
upload_response.raise_for_status()

print("Document modified and uploaded successfully.")
