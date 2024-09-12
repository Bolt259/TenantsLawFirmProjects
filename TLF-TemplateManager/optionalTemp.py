import json
from pathlib import Path
from docx import Document

def load_replacements(replacements_file: str) -> dict:
    """Load replacements from a JSON file."""
    with open(replacements_file, 'r') as file:
        return json.load(file)

def prompt_for_replacements(replacements: dict) -> dict:
    """Prompt the user for values for each replacement key."""
    updated_replacements = {}
    for key in replacements.keys():
        user_input = input(f"Enter value for '{key}': ")
        updated_replacements[key] = user_input
    return updated_replacements

def template_select(templates_path: Path) -> list:
    """Allow the user to select which templates to process."""
    templates = {
        "1": "FROGS",
        "2": "RFA",
        "3": "RFP",
        "4": "SPROG"
    }

    # Display checklist
    print("Select the templates you want to reproduce by entering the corresponding numbers (e.g., 1 3):")
    for num, template_name in templates.items():
        print(f"{num}: {template_name}")

    # User input for template selection
    user_input = input("Your choice: ").split()
    
    # Filter valid selections and build the list of files to process
    selected_templates = [templates_path / f"{templates[num]}.docx" for num in user_input if num in templates]
    
    if not selected_templates:
        print("No valid templates selected. Exiting.")
    
    return selected_templates

def replace_text_in_docx(doc: Document, replacements: dict) -> None:
    """Replace text in a .docx file according to the replacements dictionary."""
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # Replace text in tables if needed
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

def process_template(template_file: Path, replacements: dict, output_dir: Path) -> None:
    """Process a single template file by performing replacements and saving the output."""
    doc = Document(template_file)
    replace_text_in_docx(doc, replacements)
    output_file = output_dir / template_file.name
    doc.save(output_file)

def main(templates_dir: str, replacements_file: str, output_dir: str) -> None:
    """Main function to process templates based on user input and replacements."""
    templates_path = Path(templates_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    replacements = load_replacements(replacements_file)
    updated_replacements = prompt_for_replacements(replacements)
    selected_templates = template_select(templates_path)

    for template_file in selected_templates:
        process_template(template_file, updated_replacements, output_path)
        print(f"Processed {template_file.name} and saved to {output_path}")

if __name__ == "__main__":
    templates_dir = input("Enter the path to your templates directory: ")
    replacements_file = input("Enter the path to your replacements.json file: ")
    output_dir = input("Enter the path to your output directory: ")

    main(templates_dir, replacements_file, output_dir)
